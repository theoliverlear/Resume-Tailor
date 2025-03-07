from docx import Document

from src.ai.doc_ai import call_open_ai

document = Document("Test_Resume.docx")

def get_editable_skills():
    for table in document.tables:
        for row in table.rows:
            for index, cell in enumerate(row.cells):
                if "Skills" in cell.text:
                    relevant_text = cell.text.split("Skills")[1]
                    skills = relevant_text.split('\n')
                    return list(skills[1:])

default_skills = get_editable_skills()

def get_editable_skills_cells():
    for table in document.tables:
        for row in table.rows:
            for index, cell in enumerate(row.cells):
                if "Skills" in cell.text:
                    return cell

def get_skill_cell():
    for table in document.tables:
        for row in table.rows:
            for index, cell in enumerate(row.cells):
                if "Skills" in cell.text:
                    return cell

def inject_skills(cell, new_skills):
    for index, skill in enumerate(default_skills):
        if len(new_skills) > index:
            replace_text_preserving_format(cell,
                                           default_skills[index],
                                           new_skills[index])

def replace_text_preserving_format(cell, old_text, new_text):
    found_skills: bool = False
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            skills_detected: bool = not found_skills and "Skills" in run.text
            if skills_detected:
                found_skills = True
                continue
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)


def main():
    cell = get_editable_skills_cells()
    inject_skills(cell, call_open_ai())
    document.save("Updated_Resume.docx")

if __name__ == "__main__":
    main()
